"""
Advanced Document Processor with Enhanced Accuracy
Improved text extraction, structure detection, and content analysis
"""

import asyncio
import aiofiles
import httpx
import logging
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import tempfile
import os
import re
from dataclasses import dataclass

# Enhanced document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
import email
import json

logger = logging.getLogger(__name__)

@dataclass
class DocumentSection:
    """Structured representation of document sections"""
    title: str
    content: str
    level: int
    section_type: str
    confidence: float
    metadata: Dict[str, Any]

class AdvancedDocumentProcessor:
    """Enhanced document processor with improved accuracy"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.eml', '.html', '.htm']
        self.temp_dir = tempfile.mkdtemp()
        
        # Enhanced patterns for better structure detection
        self.section_patterns = {
            'heading': [
                r'^[A-Z][A-Z\s]{3,}:?\s*$',  # ALL CAPS headings
                r'^\d+\.\s+[A-Z][^.]*$',      # Numbered sections
                r'^[IVX]+\.\s+[A-Z][^.]*$',   # Roman numerals
                r'^[A-Z]\.\s+[A-Z][^.]*$',    # Letter sections
            ],
            'definition': [
                r'([A-Z][a-z\s]+)\s+means\s+',
                r'([A-Z][a-z\s]+)\s+shall\s+mean\s+',
                r'([A-Z][a-z\s]+)\s+is\s+defined\s+as\s+',
            ],
            'clause': [
                r'^\d+\.\d+\s+',              # Numbered clauses
                r'^[a-z]\)\s+',               # Lettered clauses
                r'^\([a-z]\)\s+',             # Parenthetical clauses
            ],
            'amount': [
                r'INR\s*[\d,]+',
                r'Rs\.?\s*[\d,]+',
                r'\$\s*[\d,]+',
                r'USD\s*[\d,]+',
            ],
            'percentage': [
                r'\d+(?:\.\d+)?%',
                r'\d+(?:\.\d+)?\s*percent',
            ],
            'time_period': [
                r'\d+\s*(?:days?|months?|years?)',
                r'(?:thirty|sixty|ninety)\s*(?:days?|months?)',
                r'(?:one|two|three|four|five)\s*(?:days?|months?|years?)',
            ]
        }
        
    async def process_document(self, document_url_or_path: str) -> Dict[str, Any]:
        """
        Enhanced document processing with improved accuracy
        """
        try:
            # Determine if it's a URL or local path
            if document_url_or_path.startswith(('http://', 'https://')):
                file_path = await self._download_document_enhanced(document_url_or_path)
            else:
                file_path = Path(document_url_or_path)
            
            # Process based on file extension with enhanced methods
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                content = await self._process_pdf_enhanced(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = await self._process_docx_enhanced(file_path)
            elif file_extension == '.txt':
                content = await self._process_txt_enhanced(file_path)
            elif file_extension == '.eml':
                content = await self._process_email_enhanced(file_path)
            elif file_extension in ['.html', '.htm']:
                content = await self._process_html_enhanced(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Enhanced post-processing
            content = await self._enhance_content_structure(content)
            content = await self._extract_key_entities(content)
            content = await self._improve_text_quality(content)
            
            # Add enhanced metadata
            content['metadata'].update({
                'processing_version': '2.0',
                'accuracy_enhancements': True,
                'structure_analysis': True,
                'entity_extraction': True
            })
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing document {document_url_or_path}: {str(e)}")
            raise
    
    async def _download_document_enhanced(self, url: str) -> Path:
        """Enhanced document download with better error handling"""
        try:
            timeout = httpx.Timeout(60.0, connect=30.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                # Add headers to mimic browser request
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': 'application/pdf,application/msword,text/html,*/*',
                    'Accept-Language': 'en-US,en;q=0.9',
                }
                
                response = await client.get(url, headers=headers)
                response.raise_for_status()
                
                # Better content type detection
                content_type = response.headers.get('content-type', '').lower()
                if 'pdf' in content_type:
                    extension = '.pdf'
                elif 'word' in content_type or 'document' in content_type:
                    extension = '.docx'
                elif 'html' in content_type:
                    extension = '.html'
                else:
                    # Try to get from URL or default to PDF
                    extension = Path(url).suffix or '.pdf'
                
                # Save with better naming
                temp_file = Path(self.temp_dir) / f"doc_{hash(url) % 10000}{extension}"
                
                async with aiofiles.open(temp_file, 'wb') as f:
                    await f.write(response.content)
                
                logger.info(f"Downloaded document: {temp_file} ({len(response.content)} bytes)")
                return temp_file
                
        except Exception as e:
            logger.error(f"Error downloading document from {url}: {str(e)}")
            raise
    
    async def _process_pdf_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced PDF processing with better text extraction"""
        content = {
            'text': '',
            'pages': [],
            'tables': [],
            'sections': {},
            'structured_content': [],
            'key_phrases': [],
            'entities': {}
        }
        
        try:
            # Primary extraction with pdfplumber (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                
                for page_num, page in enumerate(pdf.pages):
                    # Enhanced text extraction
                    page_text = self._extract_text_with_layout(page)
                    
                    if page_text:
                        content['pages'].append({
                            'page_number': page_num + 1,
                            'text': page_text,
                            'word_count': len(page_text.split()),
                            'has_tables': bool(page.extract_tables())
                        })
                        full_text += page_text + '\n\n'
                    
                    # Enhanced table extraction
                    tables = page.extract_tables()
                    if tables:
                        for table_idx, table in enumerate(tables):
                            processed_table = self._process_table_enhanced(table, page_num + 1, table_idx)
                            content['tables'].append(processed_table)
                
                content['text'] = full_text
            
            # Fallback with PyPDF2 if pdfplumber fails
            if not content['text'].strip():
                content = await self._fallback_pdf_processing(file_path, content)
            
            # Enhanced structure detection
            content['sections'] = self._detect_document_sections_enhanced(content['text'])
            content['structured_content'] = self._create_structured_content(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            # Final fallback
            content = await self._fallback_pdf_processing(file_path, content)
        
        return content
    
    def _extract_text_with_layout(self, page) -> str:
        """Extract text while preserving layout information"""
        try:
            # Try to maintain layout structure
            text = page.extract_text(layout=True)
            if not text:
                # Fallback to regular extraction
                text = page.extract_text()
            
            # Clean up the text while preserving structure
            if text:
                # Normalize whitespace but preserve line breaks
                lines = text.split('\n')
                cleaned_lines = []
                
                for line in lines:
                    line = line.strip()
                    if line:
                        # Remove excessive spaces but keep structure
                        line = re.sub(r' {3,}', '  ', line)
                        cleaned_lines.append(line)
                
                return '\n'.join(cleaned_lines)
            
            return ""
            
        except Exception as e:
            logger.warning(f"Layout extraction failed: {e}")
            return page.extract_text() or ""
    
    def _process_table_enhanced(self, table: List[List[str]], page_num: int, table_idx: int) -> Dict[str, Any]:
        """Enhanced table processing with better structure detection"""
        if not table or not table[0]:
            return {}
        
        # Clean and structure the table
        cleaned_table = []
        for row in table:
            cleaned_row = [cell.strip() if cell else "" for cell in row]
            if any(cleaned_row):  # Only add non-empty rows
                cleaned_table.append(cleaned_row)
        
        if not cleaned_table:
            return {}
        
        # Detect headers and structure
        headers = cleaned_table[0] if cleaned_table else []
        data_rows = cleaned_table[1:] if len(cleaned_table) > 1 else []
        
        # Convert to searchable text
        table_text = self._table_to_enhanced_text(headers, data_rows)
        
        return {
            'page': page_num,
            'table_id': table_idx,
            'headers': headers,
            'data': data_rows,
            'text_representation': table_text,
            'row_count': len(data_rows),
            'column_count': len(headers) if headers else 0
        }
    
    def _table_to_enhanced_text(self, headers: List[str], data_rows: List[List[str]]) -> str:
        """Convert table to enhanced searchable text"""
        text_parts = []
        
        if headers and data_rows:
            for row in data_rows:
                row_texts = []
                for i, cell in enumerate(row):
                    if i < len(headers) and cell.strip():
                        header = headers[i].strip()
                        if header:
                            row_texts.append(f"{header}: {cell.strip()}")
                        else:
                            row_texts.append(cell.strip())
                
                if row_texts:
                    text_parts.append("; ".join(row_texts))
        
        return ". ".join(text_parts)
    
    async def _fallback_pdf_processing(self, file_path: Path, content: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback PDF processing using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content['pages'].append({
                                'page_number': page_num + 1,
                                'text': page_text,
                                'extraction_method': 'PyPDF2_fallback'
                            })
                            full_text += page_text + '\n'
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {e}")
                
                content['text'] = full_text
                
        except Exception as e:
            logger.error(f"Fallback PDF processing failed: {str(e)}")
            content['text'] = "Error: Could not extract text from PDF"
        
        return content
    
    def _detect_document_sections_enhanced(self, text: str) -> Dict[str, DocumentSection]:
        """Enhanced section detection with confidence scoring"""
        sections = {}
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            section_info = self._analyze_line_as_header(line, line_num, lines)
            
            if section_info['is_header']:
                # Save previous section
                if current_section and current_content:
                    sections[current_section.title] = DocumentSection(
                        title=current_section.title,
                        content='\n'.join(current_content),
                        level=current_section.level,
                        section_type=current_section.section_type,
                        confidence=current_section.confidence,
                        metadata=current_section.metadata
                    )
                
                # Start new section
                current_section = DocumentSection(
                    title=line,
                    content="",
                    level=section_info['level'],
                    section_type=section_info['type'],
                    confidence=section_info['confidence'],
                    metadata={'line_number': line_num}
                )
                current_content = []
            else:
                if current_section:
                    current_content.append(line)
                else:
                    # Content before first header
                    if 'PREAMBLE' not in sections:
                        sections['PREAMBLE'] = DocumentSection(
                            title='PREAMBLE',
                            content=line,
                            level=0,
                            section_type='preamble',
                            confidence=0.8,
                            metadata={'line_number': line_num}
                        )
                    else:
                        sections['PREAMBLE'].content += '\n' + line
        
        # Save final section
        if current_section and current_content:
            sections[current_section.title] = DocumentSection(
                title=current_section.title,
                content='\n'.join(current_content),
                level=current_section.level,
                section_type=current_section.section_type,
                confidence=current_section.confidence,
                metadata=current_section.metadata
            )
        
        return {k: v.__dict__ for k, v in sections.items()}
    
    def _analyze_line_as_header(self, line: str, line_num: int, all_lines: List[str]) -> Dict[str, Any]:
        """Analyze if a line is likely a section header"""
        confidence = 0.0
        header_type = 'unknown'
        level = 0
        
        # Check various header patterns
        for pattern_type, patterns in self.section_patterns.items():
            if pattern_type == 'heading':
                for pattern in patterns:
                    if re.match(pattern, line):
                        confidence += 0.3
                        header_type = 'heading'
                        
                        # Determine level based on pattern
                        if re.match(r'^\d+\.\s+', line):
                            level = 1
                        elif re.match(r'^\d+\.\d+\s+', line):
                            level = 2
                        elif line.isupper():
                            level = 1
        
        # Additional heuristics
        if len(line) < 100 and ':' in line and not line.endswith('.'):
            confidence += 0.2
            header_type = 'section_header'
        
        if line.isupper() and len(line.split()) <= 8:
            confidence += 0.3
            level = max(level, 1)
        
        # Context analysis
        if line_num > 0 and line_num < len(all_lines) - 1:
            prev_line = all_lines[line_num - 1].strip()
            next_line = all_lines[line_num + 1].strip()
            
            # Empty lines around suggest header
            if not prev_line and not next_line:
                confidence += 0.2
            elif not prev_line or not next_line:
                confidence += 0.1
        
        return {
            'is_header': confidence > 0.4,
            'confidence': min(confidence, 1.0),
            'type': header_type,
            'level': level
        }
    
    def _create_structured_content(self, text: str) -> List[Dict[str, Any]]:
        """Create structured content blocks for better processing"""
        structured_blocks = []
        paragraphs = text.split('\n\n')
        
        for para_num, paragraph in enumerate(paragraphs):
            if not paragraph.strip():
                continue
            
            block = {
                'id': para_num,
                'text': paragraph.strip(),
                'type': self._classify_content_block(paragraph),
                'entities': self._extract_entities_from_block(paragraph),
                'key_phrases': self._extract_key_phrases_from_block(paragraph),
                'word_count': len(paragraph.split())
            }
            
            structured_blocks.append(block)
        
        return structured_blocks
    
    def _classify_content_block(self, text: str) -> str:
        """Classify content blocks by type"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['definition', 'means', 'shall mean', 'defined as']):
            return 'definition'
        elif any(word in text_lower for word in ['exclusion', 'excluded', 'not covered', 'limitation']):
            return 'exclusion'
        elif any(word in text_lower for word in ['coverage', 'covered', 'benefit', 'entitled']):
            return 'coverage'
        elif any(word in text_lower for word in ['procedure', 'process', 'steps', 'how to']):
            return 'procedure'
        elif re.search(r'\d+(?:\.\d+)?%|\$[\d,]+|inr\s*[\d,]+', text_lower):
            return 'financial'
        elif any(word in text_lower for word in ['waiting period', 'grace period', 'time limit']):
            return 'temporal'
        else:
            return 'general'
    
    def _extract_entities_from_block(self, text: str) -> Dict[str, List[str]]:
        """Extract entities from a content block"""
        entities = {
            'amounts': [],
            'percentages': [],
            'time_periods': [],
            'medical_terms': [],
            'legal_terms': []
        }
        
        # Extract amounts
        amount_patterns = self.section_patterns['amount']
        for pattern in amount_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['amounts'].extend(matches)
        
        # Extract percentages
        percentage_patterns = self.section_patterns['percentage']
        for pattern in percentage_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['percentages'].extend(matches)
        
        # Extract time periods
        time_patterns = self.section_patterns['time_period']
        for pattern in time_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['time_periods'].extend(matches)
        
        # Medical terms (basic list)
        medical_terms = [
            'surgery', 'treatment', 'diagnosis', 'procedure', 'therapy',
            'hospitalization', 'consultation', 'medication', 'prescription'
        ]
        for term in medical_terms:
            if term in text.lower():
                entities['medical_terms'].append(term)
        
        # Legal terms (basic list)
        legal_terms = [
            'clause', 'provision', 'terms', 'conditions', 'agreement',
            'contract', 'liability', 'obligation', 'rights', 'duties'
        ]
        for term in legal_terms:
            if term in text.lower():
                entities['legal_terms'].append(term)
        
        return entities
    
    def _extract_key_phrases_from_block(self, text: str) -> List[str]:
        """Extract key phrases from content block"""
        # Simple key phrase extraction
        sentences = re.split(r'[.!?]+', text)
        key_phrases = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 20 and len(sentence) < 200:
                # Look for important phrases
                if any(word in sentence.lower() for word in [
                    'covered', 'excluded', 'required', 'must', 'shall',
                    'entitled', 'eligible', 'applicable', 'subject to'
                ]):
                    key_phrases.append(sentence)
        
        return key_phrases[:5]  # Top 5 key phrases
    
    async def _enhance_content_structure(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Enhance content structure for better processing"""
        # Add cross-references between sections
        if 'sections' in content:
            content['section_relationships'] = self._analyze_section_relationships(content['sections'])
        
        # Add content hierarchy
        if 'structured_content' in content:
            content['content_hierarchy'] = self._build_content_hierarchy(content['structured_content'])
        
        return content
    
    def _analyze_section_relationships(self, sections: Dict[str, Any]) -> Dict[str, List[str]]:
        """Analyze relationships between sections"""
        relationships = {}
        
        for section_name, section_data in sections.items():
            related_sections = []
            section_content = section_data.get('content', '').lower()
            
            # Find references to other sections
            for other_section, other_data in sections.items():
                if other_section != section_name:
                    other_title = other_section.lower()
                    if other_title in section_content or any(
                        word in section_content for word in other_title.split()[:3]
                    ):
                        related_sections.append(other_section)
            
            relationships[section_name] = related_sections
        
        return relationships
    
    def _build_content_hierarchy(self, structured_content: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Build content hierarchy for better navigation"""
        hierarchy = {
            'definitions': [],
            'coverage_items': [],
            'exclusions': [],
            'procedures': [],
            'financial_info': [],
            'temporal_info': []
        }
        
        for block in structured_content:
            block_type = block.get('type', 'general')
            
            if block_type == 'definition':
                hierarchy['definitions'].append(block['id'])
            elif block_type == 'coverage':
                hierarchy['coverage_items'].append(block['id'])
            elif block_type == 'exclusion':
                hierarchy['exclusions'].append(block['id'])
            elif block_type == 'procedure':
                hierarchy['procedures'].append(block['id'])
            elif block_type == 'financial':
                hierarchy['financial_info'].append(block['id'])
            elif block_type == 'temporal':
                hierarchy['temporal_info'].append(block['id'])
        
        return hierarchy
    
    async def _extract_key_entities(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Extract key entities from the entire document"""
        text = content.get('text', '')
        
        entities = {
            'companies': [],
            'amounts': [],
            'percentages': [],
            'time_periods': [],
            'locations': [],
            'contact_info': [],
            'policy_numbers': [],
            'dates': []
        }
        
        # Company names
        company_patterns = [
            r'([A-Z][a-z]+\s+Insurance\s+Company)',
            r'([A-Z][a-z]+\s+Insurance\s+Co\.?\s*Ltd\.?)',
            r'([A-Z][a-z]+\s+Life\s+Insurance)',
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            entities['companies'].extend(matches)
        
        # Contact information
        phone_pattern = r'(\+?\d{1,3}[-.\s]?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4})'
        email_pattern = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        
        entities['contact_info'].extend(re.findall(phone_pattern, text))
        entities['contact_info'].extend(re.findall(email_pattern, text))
        
        # Policy numbers
        policy_patterns = [
            r'(Policy\s+No\.?\s*:?\s*([A-Z0-9/-]+))',
            r'(UIN\s*:?\s*([A-Z0-9]+))',
            r'(Registration\s+No\.?\s*:?\s*(\d+))'
        ]
        
        for pattern in policy_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['policy_numbers'].extend([match[1] if isinstance(match, tuple) else match for match in matches])
        
        # Dates
        date_patterns = [
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})',
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['dates'].extend(matches)
        
        content['extracted_entities'] = entities
        return content
    
    async def _improve_text_quality(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Improve text quality through cleaning and normalization"""
        text = content.get('text', '')
        
        # Clean up common OCR errors and formatting issues
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
        text = re.sub(r'(\d)([A-Z])', r'\1 \2', text)  # Add space between number and letter
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Add space after punctuation
        
        # Fix common OCR substitutions
        ocr_fixes = {
            'rn': 'm',
            '0': 'O',  # In context where O makes more sense
            '1': 'l',  # In context where l makes more sense
        }
        
        # Apply fixes contextually (basic implementation)
        for wrong, correct in ocr_fixes.items():
            # Only apply in specific contexts to avoid over-correction
            pass  # Implement contextual fixes as needed
        
        content['text'] = text
        content['text_quality_improved'] = True
        
        return content
    
    async def _process_docx_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced DOCX processing"""
        # Implementation similar to PDF but for DOCX
        # This would include better table extraction, style analysis, etc.
        return await self._process_docx(file_path)  # Fallback to original for now
    
    async def _process_txt_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced TXT processing"""
        # Implementation for enhanced text processing
        return await self._process_txt(file_path)  # Fallback to original for now
    
    async def _process_email_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced email processing"""
        # Implementation for enhanced email processing
        return await self._process_email(file_path)  # Fallback to original for now
    
    async def _process_html_enhanced(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced HTML processing"""
        # Implementation for enhanced HTML processing
        return await self._process_html(file_path)  # Fallback to original for now
    
    # Include original methods as fallbacks
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Original DOCX processing as fallback"""
        content = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'sections': {}
        }
        
        try:
            doc = Document(file_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    content['text'] += para.text + '\n'
            
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'table_number': table_num + 1,
                    'data': table_data
                })
            
            content['sections'] = self._detect_document_sections_enhanced(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
        
        return content
    
    async def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Original TXT processing as fallback"""
        content = {
            'text': '',
            'sections': {}
        }
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
                content['text'] = text
                content['sections'] = self._detect_document_sections_enhanced(text)
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        text = await f.read()
                        content['text'] = text
                        content['sections'] = self._detect_document_sections_enhanced(text)
                        break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Could not decode text file {file_path}")
        
        return content
    
    async def _process_email(self, file_path: Path) -> Dict[str, Any]:
        """Original email processing as fallback"""
        content = {
            'text': '',
            'subject': '',
            'sender': '',
            'recipients': [],
            'date': '',
            'sections': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                msg = email.message_from_file(f)
            
            content['subject'] = msg.get('Subject', '')
            content['sender'] = msg.get('From', '')
            content['recipients'] = msg.get_all('To', [])
            content['date'] = msg.get('Date', '')
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        content['text'] += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                content['text'] = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            content['sections'] = self._detect_document_sections_enhanced(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing email {file_path}: {str(e)}")
            raise
        
        return content
    
    async def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Original HTML processing as fallback"""
        content = {
            'text': '',
            'sections': {}
        }
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                html_content = await f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            content['text'] = soup.get_text(separator='\n', strip=True)
            content['sections'] = self._detect_document_sections_enhanced(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {str(e)}")
            raise
        
        return content