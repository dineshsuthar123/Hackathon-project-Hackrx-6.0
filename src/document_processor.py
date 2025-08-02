"""
Document Processor Module
Handles processing of PDFs, DOCX, and email documents
"""

import asyncio
import aiofiles
import httpx
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import tempfile
import os

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
import email
from email.mime.text import MIMEText

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes various document formats and extracts structured content"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.eml', '.html']
        self.temp_dir = tempfile.mkdtemp()
        
    async def process_document(self, document_url_or_path: str) -> Dict[str, Any]:
        """
        Main document processing method
        
        Args:
            document_url_or_path: URL or file path to the document
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            # Determine if it's a URL or local path
            if document_url_or_path.startswith(('http://', 'https://')):
                file_path = await self._download_document(document_url_or_path)
            else:
                file_path = Path(document_url_or_path)
            
            # Process based on file extension
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                content = await self._process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = await self._process_docx(file_path)
            elif file_extension == '.txt':
                content = await self._process_txt(file_path)
            elif file_extension == '.eml':
                content = await self._process_email(file_path)
            elif file_extension == '.html':
                content = await self._process_html(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Add metadata
            content['metadata'] = {
                'source': document_url_or_path,
                'file_type': file_extension,
                'processed_at': str(asyncio.get_event_loop().time()),
                'size_kb': file_path.stat().st_size / 1024 if file_path.exists() else 0
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing document {document_url_or_path}: {str(e)}")
            raise
    
    async def _download_document(self, url: str) -> Path:
        """Download document from URL to temporary file"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Determine file extension from URL or content type
                content_type = response.headers.get('content-type', '')
                if 'pdf' in content_type:
                    extension = '.pdf'
                elif 'word' in content_type or 'document' in content_type:
                    extension = '.docx'
                else:
                    # Try to get from URL
                    extension = Path(url).suffix or '.pdf'
                
                # Save to temporary file
                temp_file = Path(self.temp_dir) / f"downloaded_doc{extension}"
                
                async with aiofiles.open(temp_file, 'wb') as f:
                    await f.write(response.content)
                
                logger.info(f"Downloaded document: {temp_file}")
                return temp_file
                
        except Exception as e:
            logger.error(f"Error downloading document from {url}: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and structure from PDF files"""
        content = {
            'text': '',
            'pages': [],
            'tables': [],
            'sections': {}
        }
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ''
                    content['pages'].append({
                        'page_number': page_num + 1,
                        'text': page_text
                    })
                    content['text'] += page_text + '\n'
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            content['tables'].append({
                                'page': page_num + 1,
                                'data': table
                            })
            
            # Structure detection (simple approach)
            content['sections'] = self._detect_document_sections(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        content['pages'].append({
                            'page_number': page_num + 1,
                            'text': page_text
                        })
                        content['text'] += page_text + '\n'
            except Exception as fallback_error:
                logger.error(f"Fallback PDF processing failed: {str(fallback_error)}")
                raise
        
        return content
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and structure from DOCX files"""
        content = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'sections': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    content['text'] += para.text + '\n'
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'table_number': table_num + 1,
                    'data': table_data
                })
            
            # Structure detection
            content['sections'] = self._detect_document_sections(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
        
        return content
    
    async def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text files"""
        content = {
            'text': '',
            'sections': {}
        }
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
                content['text'] = text
                content['sections'] = self._detect_document_sections(text)
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                        text = await f.read()
                        content['text'] = text
                        content['sections'] = self._detect_document_sections(text)
                        break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Could not decode text file {file_path}")
        
        return content
    
    async def _process_email(self, file_path: Path) -> Dict[str, Any]:
        """Process email files (.eml)"""
        content = {
            'text': '',
            'subject': '',
            'sender': '',
            'recipients': [],
            'date': '',
            'sections': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                msg = email.message_from_file(f)
            
            content['subject'] = msg.get('Subject', '')
            content['sender'] = msg.get('From', '')
            content['recipients'] = msg.get_all('To', [])
            content['date'] = msg.get('Date', '')
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        content['text'] += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                content['text'] = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            content['sections'] = self._detect_document_sections(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing email {file_path}: {str(e)}")
            raise
        
        return content
    
    async def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Process HTML files"""
        content = {
            'text': '',
            'sections': {}
        }
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                html_content = await f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            content['text'] = soup.get_text(separator='\n', strip=True)
            content['sections'] = self._detect_document_sections(content['text'])
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {str(e)}")
            raise
        
        return content
    
    def _detect_document_sections(self, text: str) -> Dict[str, str]:
        """
        Simple document section detection based on common patterns
        This can be enhanced with more sophisticated NLP techniques
        """
        sections = {}
        
        # Common section headers
        section_patterns = [
            'INTRODUCTION', 'BACKGROUND', 'OVERVIEW', 'SUMMARY',
            'COVERAGE', 'BENEFITS', 'EXCLUSIONS', 'CONDITIONS',
            'DEFINITIONS', 'TERMS', 'POLICY', 'CLAIMS',
            'PREMIUM', 'DEDUCTIBLE', 'WAITING PERIOD'
        ]
        
        lines = text.split('\n')
        current_section = 'GENERAL'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_section_header = False
            for pattern in section_patterns:
                if pattern.lower() in line.lower() and len(line) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = line.upper()
                    current_content = []
                    is_section_header = True
                    break
            
            if not is_section_header:
                current_content.append(line)
        
        # Save final section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def __del__(self):
        """Cleanup temporary files"""
        import shutil
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception:
            pass
